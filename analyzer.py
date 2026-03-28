"""
analyzer.py — Professional ATS Resume Scoring Engine
======================================================
Scoring methodology based on how real ATS tools work:
  - Jobscan, Resume Worded, VMock, Enhancv
  - Weighted across 8 real criteria
  - Keyword matching against 200+ industry terms
  - Smart section detection
  - Quantification analysis
  - Readability & format checks
"""

import re
import io as _io
from collections import Counter

# ── PDF ───────────────────────────────────────────────────────
try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# ── DOCX ─────────────────────────────────────────────────────
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ── OCR ──────────────────────────────────────────────────────
try:
    import pytesseract
    from PIL import Image
    from pdf2image import convert_from_bytes
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False


# ═══════════════════════════════════════════════════════════════
# KEYWORD BANKS  — 200+ terms across industries
# ═══════════════════════════════════════════════════════════════

# Action verbs that top ATS tools look for
ACTION_VERBS = [
    "achieved","administered","analyzed","architected","automated","built",
    "collaborated","created","coordinated","delivered","deployed","designed",
    "developed","directed","drove","engineered","established","evaluated",
    "executed","generated","implemented","improved","increased","launched",
    "led","managed","mentored","optimized","orchestrated","oversaw",
    "planned","produced","reduced","resolved","scaled","spearheaded",
    "streamlined","supervised","trained","transformed","developed","utilized",
    "accelerated","awarded","championed","consolidated","constructed",
    "contributed","cultivated","customized","decreased","defined","demonstrated",
    "enhanced","expanded","facilitated","formulated","grew","guided","handled",
    "identified","initiated","integrated","introduced","investigated","negotiated",
    "operated","organized","performed","prepared","presented","prioritized",
    "processed","programmed","proposed","provided","published","recruited",
    "restructured","reviewed","secured","simplified","solved","supported",
    "tested","upgraded","validated","wrote"
]

# Weak filler phrases real ATS tools penalise
WEAK_PHRASES = [
    "responsible for","duties included","tasks included","worked on",
    "helped with","assisted in","participated in","involved in",
    "was part of","tried to","attempted to","helped to",
    "contributed to the team","was responsible","job duties"
]

# ── Section headers ───────────────────────────────────────────
SECTION_PATTERNS = {
    "Contact Info":    r'(email|phone|mobile|tel|linkedin|github|address|location|city|contact)',
    "Summary":         r'(summary|objective|profile|about me|overview|professional summary|career objective|personal statement)',
    "Experience":      r'(experience|employment|work history|career|positions? held|professional experience|internship|work experience)',
    "Education":       r'(education|academic|qualification|degree|university|college|school|bachelor|master|phd|mba|diploma)',
    "Skills":          r'(skill|technolog|tools?|languages?|frameworks?|competenc|expertise|proficienc|tech stack|technical)',
    "Projects":        r'(project|portfolio|open.?source|side project|personal project)',
    "Certifications":  r'(certif|license|accredit|credential|aws certified|google certified|microsoft certified)',
    "Achievements":    r'(achievement|award|honor|recognition|publication|patent|speaker|accomplishment)',
}

# ── Universal keyword bank (200+ terms) ──────────────────────
UNIVERSAL_KEYWORDS = {
    # Programming languages
    "python","java","javascript","typescript","c++","c#","go","rust",
    "swift","kotlin","ruby","php","scala","r","matlab","perl","bash","shell",

    # Web & frameworks
    "react","angular","vue","node","express","django","flask","fastapi",
    "spring","laravel","rails","next.js","nuxt","svelte","html","css",
    "tailwind","bootstrap","jquery","webpack","vite",

    # Databases
    "sql","mysql","postgresql","mongodb","redis","cassandra","dynamodb",
    "oracle","sqlite","elasticsearch","snowflake","bigquery","redshift",

    # Cloud & DevOps
    "aws","azure","gcp","docker","kubernetes","terraform","ansible",
    "jenkins","github actions","ci/cd","linux","nginx","apache",
    "cloudformation","lambda","ec2","s3","heroku","vercel",

    # Data science & ML
    "machine learning","deep learning","nlp","computer vision","tensorflow",
    "pytorch","keras","scikit-learn","pandas","numpy","matplotlib","seaborn",
    "xgboost","lightgbm","spark","hadoop","airflow","mlflow","mlops",
    "data science","data analysis","statistical modeling","predictive modeling",
    "regression","classification","clustering","neural network","transformer",
    "llm","generative ai","feature engineering","model deployment",
    "a/b testing","hypothesis testing","statistical analysis",

    # Data & BI
    "tableau","power bi","looker","dbt","etl","data pipeline","data warehouse",
    "analytics","business intelligence","excel","google analytics","mixpanel",

    # Project management & methodology
    "agile","scrum","kanban","jira","confluence","trello","asana",
    "project management","product management","roadmap","sprint","stakeholder",

    # Security & networking
    "cybersecurity","penetration testing","firewall","vpn","ssl","oauth",
    "jwt","encryption","soc","compliance","gdpr","iso 27001",

    # General tech
    "api","rest api","graphql","microservices","git","github","gitlab",
    "bitbucket","figma","adobe","photoshop","illustrator","ux","ui",
    "mobile","android","ios","react native","flutter","swift",

    # Business & finance (only universal ones — not role-specific)
    "revenue","cost reduction","roi","kpi","metrics","reporting",
    "salesforce","crm","budgeting","forecasting","p&l",

    # Marketing (only widely relevant)
    "seo","sem","content marketing","social media","email marketing",
    "market research","digital marketing","google ads","branding",

    # Healthcare (only widely relevant)
    "hipaa","clinical","patient care","healthcare","medical",

    # Soft skills (that appear as keywords)
    "leadership","communication","collaboration","teamwork","problem solving",
    "critical thinking","analytical","strategic","cross-functional",
    "stakeholder management","mentoring","coaching","decision making",
    "time management","project management","innovation","customer success",
}

# Common words to ignore in JD matching
STOPWORDS = {
    "the","a","an","and","or","but","in","on","at","to","for","of","with",
    "by","from","is","are","was","were","be","been","have","has","had",
    "do","does","did","will","would","could","should","may","might",
    "this","that","these","those","it","its","you","we","they","he","she",
    "our","their","your","his","her","as","not","no","if","so","up",
    "can","also","all","any","both","each","few","more","most","other",
    "some","such","than","then","too","very","just","about","into","over",
    "after","above","below","between","through","during","before","under",
    "again","further","once","how","what","when","where","who","which","why",
}



# ═══════════════════════════════════════════════════════════════
# ROLE-SPECIFIC KEYWORD BANKS — CLEANED & SMART
# Only relevant keywords per role. No cross-role pollution.
# Missing keywords = actually relevant to that role.
# ═══════════════════════════════════════════════════════════════

ROLE_KEYWORDS = {

    "general": set(),  # empty = use UNIVERSAL_KEYWORDS

    "data_science": {
        # Core technical skills
        "python","r","sql","machine learning","deep learning","nlp",
        "computer vision","tensorflow","pytorch","keras","scikit-learn",
        "pandas","numpy","matplotlib","seaborn","xgboost","lightgbm",
        "spark","hadoop","airflow","mlflow","mlops",
        # Analysis & modeling
        "data science","data analysis","statistical modeling",
        "predictive modeling","regression","classification","clustering",
        "neural network","transformer","llm","generative ai",
        "feature engineering","model deployment","hypothesis testing",
        "statistical analysis","probability","statistics","time series",
        # Data tools
        "tableau","power bi","looker","dbt","etl","data pipeline",
        "data warehouse","bigquery","snowflake","redshift","databricks",
        "sagemaker","jupyter","data visualization","business intelligence",
        # Cloud & DevOps basics
        "aws","azure","gcp","docker","git","linux",
        # Soft skills relevant to DS
        "communication","presentation","stakeholder management",
        "cross-functional","problem solving","critical thinking",
    },

    "frontend": {
        # Core languages
        "html","css","javascript","typescript",
        # Frameworks
        "react","angular","vue","next.js","nuxt","svelte",
        # Styling
        "tailwind","bootstrap","sass","scss","css grid","flexbox",
        # Build tools
        "webpack","vite","babel","npm","yarn",
        # Version control
        "git","github",
        # API
        "rest api","graphql","fetch api","axios",
        # Testing
        "jest","cypress","testing","unit testing",
        # UX concepts
        "accessibility","responsive design","mobile first",
        "performance optimization","seo","web vitals","pwa",
        "cross-browser compatibility","ui","ux",
        # Design tools
        "figma","design systems","component library","storybook",
        # State management
        "redux","zustand","context api",
        # Other
        "agile","scrum","jira","ci/cd","vercel","netlify","node",
    },

    "backend": {
        # Languages
        "python","java","node","go","rust","c#","php","ruby","scala",
        # Frameworks
        "django","flask","fastapi","spring","express","rails","laravel",
        # API
        "rest api","graphql","microservices","grpc","websockets",
        # Databases
        "sql","postgresql","mysql","mongodb","redis","cassandra",
        "elasticsearch","database design","orm",
        # Message queues
        "rabbitmq","kafka","message queues","event driven",
        # DevOps
        "docker","kubernetes","aws","azure","gcp","terraform",
        "ansible","ci/cd","jenkins","linux","nginx","apache",
        # Security
        "authentication","oauth","jwt","security",
        # Other
        "git","github","testing","tdd","agile","scrum",
        "scalability","caching","performance","system design",
    },

    "fullstack": {
        # Frontend
        "html","css","javascript","typescript","react","vue","angular",
        "next.js","tailwind","bootstrap","responsive design",
        # Backend
        "node","express","python","django","flask","rest api","graphql",
        # Databases
        "sql","postgresql","mongodb","redis",
        # DevOps
        "docker","aws","git","ci/cd","linux","nginx","vercel",
        # Concepts
        "agile","testing","authentication","jwt","oauth",
        "microservices","system design","performance","security",
        "webpack","vite","npm","figma","ui","ux",
        "mobile first","database design",
    },

    "devops": {
        # Containers & orchestration
        "docker","kubernetes","helm","istio","service mesh",
        # IaC
        "terraform","ansible","puppet","chef","cloudformation",
        "infrastructure as code",
        # Cloud
        "aws","azure","gcp",
        # CI/CD
        "jenkins","github actions","gitlab ci","ci/cd",
        # Monitoring
        "prometheus","grafana","elk stack","elasticsearch",
        "logstash","kibana","monitoring","alerting","logging",
        # Scripting
        "linux","bash","python","go","scripting","automation",
        # Networking
        "networking","tcp/ip","dns","load balancing","ssl","tls",
        "nginx","apache",
        # Concepts
        "site reliability","sre","devops","security","compliance",
        "incident management","disaster recovery","cost optimization",
        "git","agile",
    },

    "data_analyst": {
        # Core tools
        "sql","excel","python","r","tableau","power bi","looker",
        "google analytics","google sheets","pivot tables","vlookup",
        # Data platforms
        "bigquery","snowflake","redshift","postgresql","mysql",
        "dbt","airflow","mixpanel",
        # Analysis skills
        "data analysis","data visualization","reporting","dashboards",
        "kpi","metrics","etl","data pipeline","statistics",
        "hypothesis testing","regression","forecasting",
        "cohort analysis","funnel analysis","segmentation",
        "customer analytics","market research","business intelligence",
        # Libraries
        "pandas","numpy","matplotlib","seaborn",
        # Soft skills
        "stakeholder management","presentation","storytelling",
        "communication","cross-functional","agile","problem solving",
    },

    "product_manager": {
        # Core PM skills
        "product management","product roadmap","product strategy",
        "product vision","go-to-market","product launch","mvp",
        "feature prioritization","backlog grooming","sprint planning",
        # Methodologies
        "agile","scrum","kanban",
        # Tools
        "jira","confluence","trello","asana","figma","wireframes",
        # Research
        "user research","user interviews","usability testing",
        "customer feedback","market analysis","competitive analysis",
        # Metrics
        "kpi","okr","north star metric","product metrics","analytics",
        "google analytics","mixpanel","data-driven","a/b testing",
        # Requirements
        "business requirements","technical requirements",
        # Soft skills
        "stakeholder management","cross-functional","leadership",
        "communication","presentation","sql","excel","roadmap",
    },

    "marketing": {
        # Digital marketing
        "seo","sem","ppc","google ads","facebook ads","meta ads",
        "digital marketing","inbound marketing","growth hacking",
        # Content & email
        "content marketing","email marketing","copywriting",
        "content creation","social media marketing","social media",
        # Tools
        "hubspot","marketo","salesforce","mailchimp",
        "google analytics","semrush","ahrefs",
        # Analytics
        "analytics","kpi","roi","conversion rate","ctr","cpc","cpa",
        "funnel optimization","marketing automation","ab testing",
        "keyword research","backlinks",
        # Strategy
        "brand strategy","market research","campaign management",
        "branding","market segmentation","customer acquisition",
        "customer retention","lead generation","b2b","b2c",
        # Other
        "influencer marketing","video marketing","community management",
        "pr","storytelling","communication","presentation",
    },

    "ui_ux": {
        # Core UX skills
        "ui design","ux design","user research","user testing",
        "usability testing","wireframing","prototyping",
        "information architecture","user flows","journey mapping",
        "persona","heuristic evaluation","design thinking",
        # Design tools
        "figma","sketch","adobe xd","invision","zeplin",
        # Design systems
        "design systems","component library","design tokens",
        "typography","color theory","grid systems","visual design",
        # Accessibility
        "accessibility","wcag",
        # Interaction
        "interaction design","responsive design","mobile design",
        "micro-interactions","motion design","animations",
        # Web basics
        "css","html",
        # Process
        "agile","scrum","jira","confluence","a/b testing",
        "cross-functional","stakeholder management",
        "communication","presentation","analytics","data-driven",
    },

    "cybersecurity": {
        # Core skills
        "cybersecurity","penetration testing","ethical hacking",
        "vulnerability assessment","network security","firewall",
        # Security tools
        "ids","ips","siem","soc","wireshark","metasploit",
        "burpsuite","kali linux","nmap","snort",
        # Concepts
        "threat intelligence","incident response","forensics",
        "malware analysis","reverse engineering","owasp",
        "zero trust","cloud security",
        # Protocols & standards
        "ssl","tls","vpn","oauth","saml","iam","pam",
        "gdpr","hipaa","iso 27001","nist","soc 2","compliance",
        "risk management","security audit",
        # Platforms
        "aws security","azure security","linux","python","bash",
        # Certifications
        "oscp","cissp","ceh","security+","certifications",
        # Soft skills
        "problem solving","communication","documentation",
    },

    "finance": {
        # Core skills
        "financial analysis","financial modeling","valuation","dcf",
        "budgeting","forecasting","variance analysis",
        "cash flow","balance sheet","income statement",
        # Standards
        "gaap","ifrs","compliance","regulatory","risk management",
        # Specializations
        "investment banking","equity research","portfolio management",
        "m&a","derivatives","options","bonds","equities",
        "fixed income","quantitative","basel","solvency",
        # Metrics
        "roi","npv","irr","wacc","ebitda","p&l",
        # Tools
        "excel","vba","bloomberg","python","r","sql",
        "tableau","power bi",
        # Certifications
        "cfa","cpa","frm","series 7","series 63",
        # Soft skills
        "stakeholder management","presentation","communication",
        "leadership","attention to detail","problem solving",
    },

    "hr": {
        # Core HR
        "human resources","talent acquisition","recruitment","sourcing",
        "onboarding","employee relations","performance management",
        "compensation","benefits","payroll","training",
        # Tools
        "hris","workday","bamboohr","adp","linkedin recruiter",
        "applicant tracking","ats",
        # Development
        "learning development","organizational development",
        "succession planning","workforce planning",
        # Culture
        "diversity","inclusion","belonging","dei",
        "employee engagement","culture","employer branding",
        # Process
        "job descriptions","interviews","background checks",
        "employment law","compliance","fmla","ada","eeoc",
        "labor relations","change management",
        # Analytics
        "kpi","metrics","analytics","excel","reporting","data-driven",
        # Soft skills
        "stakeholder management","communication","leadership",
        "presentation","empathy","problem solving",
    },

    "healthcare": {
        # Clinical skills
        "patient care","clinical","diagnosis","treatment planning",
        "nursing","pharmacology","anatomy","physiology",
        "vital signs","triage","assessment","documentation",
        # Systems
        "emr","ehr","epic","cerner","icd-10","cpt",
        # Compliance
        "hipaa","fda","irb","gcp","protocols","regulatory","compliance",
        # Research
        "clinical research","clinical trials","evidence-based",
        "quality improvement","infection control",
        # Care coordination
        "case management","care coordination","multidisciplinary",
        # Certifications
        "bls","acls","pals","certifications","medical terminology",
        # Soft skills
        "leadership","communication","empathy","teamwork",
        "critical thinking","problem solving","attention to detail",
    },
}

# ── Role display names for frontend ──────────────────────────
# ── Role display names for frontend ──────────────────────────
ROLE_DISPLAY = {
    "general":       "General / No specific role",
    "data_science":  "Data Science & ML",
    "frontend":      "Frontend Development",
    "backend":       "Backend Development",
    "fullstack":     "Full Stack Development",
    "devops":        "DevOps & Cloud",
    "data_analyst":  "Data Analyst",
    "product_manager":"Product Management",
    "marketing":     "Marketing & Growth",
    "ui_ux":         "UI/UX Design",
    "cybersecurity": "Cybersecurity",
    "finance":       "Finance & Banking",
    "hr":            "Human Resources",
    "healthcare":    "Healthcare & Medical",
}


class ATSAnalyzer:
    """
    Professional ATS scoring engine.
    Scores resumes across 8 criteria matching real-world ATS behavior.
    """

    # ── File extraction ───────────────────────────────────────

    def extract_pdf(self, file_obj) -> str:
        if not PDF_AVAILABLE:
            raise ImportError("Run: pip install pdfplumber")
        raw = file_obj.read()
        text = ""
        with pdfplumber.open(_io.BytesIO(raw)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
        if len(text.strip()) < 80:
            if not OCR_AVAILABLE:
                raise ImportError(
                    "Scanned PDF detected. Run: pip install pytesseract pdf2image Pillow\n"
                    "Mac: brew install tesseract"
                )
            text = self._ocr_pdf(raw)
        return text.strip()

    def _ocr_pdf(self, raw: bytes) -> str:
        pages = convert_from_bytes(raw, dpi=300)
        out = ""
        for p in pages:
            out += pytesseract.image_to_string(p.convert("L"), config="--psm 6") + "\n"
        if len(out.strip()) < 30:
            raise ValueError("OCR could not read this scanned PDF. Ensure scan quality is high.")
        return out.strip()

    def extract_docx(self, file_obj) -> str:
        if not DOCX_AVAILABLE:
            raise ImportError("Run: pip install python-docx")
        doc = Document(file_obj)
        seen, lines = set(), []
        def add(t):
            t = t.strip()
            if t and t not in seen:
                seen.add(t); lines.append(t)
        for p in doc.paragraphs: add(p.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs: add(p.text)
        text = "\n".join(lines)
        if len(text.strip()) < 50:
            raise ValueError("Could not extract text from this DOCX file.")
        return text.strip()

    def extract_image(self, file_obj) -> str:
        if not OCR_AVAILABLE:
            raise ImportError("Run: pip install pytesseract Pillow\nMac: brew install tesseract")
        raw = file_obj.read()
        img = Image.open(_io.BytesIO(raw)).convert("L")
        text = pytesseract.image_to_string(img, config="--psm 6")
        if len(text.strip()) < 30:
            raise ValueError("OCR could not read this image. Ensure it is clear and high resolution.")
        return text.strip()

    # ── Main analysis entry point ─────────────────────────────

    def analyze(self, resume_text: str, job_description: str = "", role: str = "general") -> dict:
        lower = resume_text.lower()
        tokens = self._tokenize(resume_text)
        token_set = set(tokens)

        # ── Run all 8 scoring modules ─────────────────────────
        r_contact  = self._score_contact(resume_text, lower)
        r_sections = self._score_sections(lower)
        r_keywords = self._score_keywords(token_set, lower, role)
        r_writing  = self._score_writing(lower)
        r_quant    = self._score_quantification(resume_text)
        r_format   = self._score_format(resume_text)
        r_soft     = self._score_soft_skills(lower)
        r_length   = self._score_length(resume_text)

        # ── Weights (total = 100) ─────────────────────────────
        # Based on Jobscan's published scoring methodology
        weights = {
            "contact":   8,   # contact completeness
            "sections":  12,  # resume structure
            "keywords":  22,  # hard skill keyword match
            "writing":   15,  # action verbs, no filler
            "quant":     18,  # measurable achievements (biggest differentiator)
            "format":    8,   # length, readability, no tables/graphics issue
            "soft":      10,  # soft skill keywords
            "length":    7,   # word count & density
        }

        raw_scores = {
            "contact":  r_contact["score"],
            "sections": r_sections["score"],
            "keywords": r_keywords["score"],
            "writing":  r_writing["score"],
            "quant":    r_quant["score"],
            "format":   r_format["score"],
            "soft":     r_soft["score"],
            "length":   r_length["score"],
        }

        # Weighted total
        total = sum((raw_scores[k] / 100) * weights[k] for k in weights)
        ats_score = round(min(99, max(10, total)))

        # ── JD matching (separate — not in base score) ────────
        jd_result = {}
        if job_description.strip():
            jd_result = self._score_jd_match(token_set, lower, job_description)
            # JD match adjusts final score slightly
            if jd_result["match_percentage"] >= 70:
                ats_score = min(99, ats_score + 3)
            elif jd_result["match_percentage"] < 30:
                ats_score = max(10, ats_score - 3)

        # ── Build score breakdown for UI ─────────────────────
        breakdown = {}
        max_breakdown = {}
        label_map = {
            "contact":   "Contact Details",
            "sections":  "Section Structure",
            "keywords":  "Technical Skills",
            "writing":   "Language & Tone",
            "quant":     "Measurable Achievements",
            "format":    "File Format & Readability",
            "soft":      "Interpersonal Skills",
            "length":    "Resume Length"
        }
        for k in weights:
            earned = round((raw_scores[k] / 100) * weights[k], 1)
            breakdown[label_map[k]] = earned
            max_breakdown[label_map[k]] = weights[k]

        # ── Improvements ──────────────────────────────────────
        improvements = self._build_improvements(raw_scores, r_contact, r_sections,
                                                 r_keywords, r_writing, r_quant,
                                                 r_format, r_soft, jd_result)

        return {
            "ats_score":          ats_score,
            "score_breakdown":    breakdown,
            "max_breakdown":      max_breakdown,
            "detected_sections":  r_sections["found"],
            "missing_sections":   r_sections["missing"],
            "found_keywords":     r_keywords["found"][:24],
            "missing_keywords":   r_keywords["missing"][:16],
            "action_verbs_found": r_writing["verbs"][:12],
            "weak_phrases_found": r_writing["weak"],
            "improvements":       improvements,
            "jd_match":           jd_result,
            "word_count":         len(resume_text.split()),
            "role":               role,
            "sentence_count":     len(re.findall(r'[.!?]+', resume_text)),
        }

    # ══════════════════════════════════════════════════════════
    # SCORING MODULES
    # ══════════════════════════════════════════════════════════

    def _score_contact(self, text: str, lower: str) -> dict:
        """
        Real ATS tools check: email, phone, LinkedIn, location.
        GitHub is bonus. Max 100, weighted by importance.
        """
        found, missing = [], []
        score = 0

        checks = [
            ("Email",    r'[\w\.\+\-]+@[\w\.\-]+\.\w{2,}',            35),
            ("Phone",    r'(\+?\d[\d\s\-\(\)\.]{6,14}\d)',             30),
            ("LinkedIn", r'linkedin\.com/in/[\w\-]+|linkedin\.com',    20),
            ("Location", r'(city|location|\b[A-Z][a-z]+,\s*[A-Z]{2}\b|\b[A-Z][a-z]+,\s*[A-Z][a-z]+)', 10),
            ("GitHub",   r'github\.com/[\w\-]+|github\.com',            5),
        ]
        for label, pattern, pts in checks:
            if re.search(pattern, text, re.IGNORECASE):
                found.append(label); score += pts
            else:
                missing.append(label)

        return {"score": min(100, score), "found": found, "missing": missing}

    def _score_sections(self, lower: str) -> dict:
        """
        Detect standard resume sections.
        Missing critical sections (Experience, Education) = heavy penalty.
        """
        found, missing = [], []
        critical = {"Experience", "Education", "Skills"}

        for section, pattern in SECTION_PATTERNS.items():
            if re.search(pattern, lower):
                found.append(section)
            else:
                missing.append(section)

        if not found:
            return {"score": 0, "found": [], "missing": list(SECTION_PATTERNS.keys())}

        base = (len(found) / len(SECTION_PATTERNS)) * 100

        # Penalty for missing critical sections
        penalty = 0
        for c in critical:
            if c not in found:
                penalty += 15

        score = max(0, base - penalty)

        # Bonus: having a summary/profile section (many resumes skip it)
        if "Summary" in found:
            score = min(100, score + 5)

        return {"score": round(score), "found": found, "missing": missing}

    def _score_keywords(self, token_set: set, lower: str, role: str = "general") -> dict:
        """
        Match resume against role-specific or universal keywords.
        SMART recommendations:
        - Found = keywords present in resume
        - Missing = only TOP PRIORITY keywords for that role, not all
        - Limits missing to 8 most important ones to avoid noise
        """
        found, missing = [], []

        # Use role-specific keywords if available, else universal
        keyword_bank = ROLE_KEYWORDS.get(role, set())
        if not keyword_bank:
            keyword_bank = UNIVERSAL_KEYWORDS

        # Priority keywords per role — shown first in recommendations
        # These are the MOST important missing keywords to show
        PRIORITY_KEYWORDS = {
            "data_science":    ["python","sql","machine learning","pandas","numpy",
                               "data analysis","statistics","git","tableau","power bi"],
            "frontend":        ["javascript","react","typescript","html","css",
                               "git","responsive design","rest api","testing","agile"],
            "backend":         ["python","sql","rest api","docker","git",
                               "postgresql","authentication","testing","linux","agile"],
            "fullstack":       ["javascript","react","node","sql","git",
                               "docker","rest api","testing","agile","postgresql"],
            "devops":          ["docker","kubernetes","aws","ci/cd","linux",
                               "terraform","git","monitoring","bash","automation"],
            "data_analyst":    ["sql","excel","tableau","python","power bi",
                               "data visualization","reporting","kpi","analytics","communication"],
            "product_manager": ["product roadmap","agile","user research","kpi",
                               "stakeholder management","jira","okr","analytics","communication","leadership"],
            "marketing":       ["seo","google analytics","content marketing","email marketing",
                               "social media","kpi","roi","campaign management","communication","analytics"],
            "ui_ux":           ["figma","user research","wireframing","prototyping",
                               "usability testing","design systems","accessibility","agile","communication","ux design"],
            "cybersecurity":   ["penetration testing","network security","siem","linux",
                               "python","compliance","risk management","incident response","firewall","owasp"],
            "finance":         ["excel","financial modeling","financial analysis","sql",
                               "forecasting","budgeting","risk management","communication","leadership","reporting"],
            "hr":              ["recruitment","talent acquisition","hris","performance management",
                               "employee relations","onboarding","compliance","communication","leadership","analytics"],
            "healthcare":      ["patient care","emr","hipaa","clinical","documentation",
                               "communication","teamwork","critical thinking","compliance","assessment"],
            "general":         ["communication","leadership","teamwork","problem solving",
                               "project management","analytical","microsoft office","presentation","agile","time management"],
        }

        for kw in sorted(keyword_bank):
            if kw in lower or kw in token_set:
                found.append(kw)
            else:
                missing.append(kw)

        # Smart missing — prioritize most important keywords first
        priority = PRIORITY_KEYWORDS.get(role, [])
        smart_missing = []

        # First add priority keywords that are missing
        for kw in priority:
            if kw in missing and kw not in smart_missing:
                smart_missing.append(kw)

        # Then add other missing keywords up to limit of 8
        for kw in missing:
            if kw not in smart_missing:
                smart_missing.append(kw)
            if len(smart_missing) >= 8:
                break

        n = len(found)
        # Scoring curve matching Jobscan behavior
        if n == 0:       score = 10
        elif n <= 5:     score = 10 + n * 8
        elif n <= 10:    score = 50 + (n - 5) * 4
        elif n <= 15:    score = 70 + (n - 10) * 2.4
        elif n <= 20:    score = 82 + (n - 15) * 1.6
        elif n <= 30:    score = 90 + (n - 20) * 0.5
        else:            score = 95 + min(3, (n - 30) * 0.15)

        return {"score": round(min(98, score)), "found": found, "missing": smart_missing}

    def _score_writing(self, lower: str) -> dict:
        """
        Real ATS tools check:
        1. Bullet points start with action verbs
        2. No passive / weak language
        3. Consistent tense
        """
        verbs_found = [v for v in ACTION_VERBS if re.search(r'\b' + v + r'\b', lower)]
        weak_found  = [w for w in WEAK_PHRASES if w in lower]

        n_verbs = len(verbs_found)
        # Scoring: 0 verbs=25, 3=60, 6=80, 10+=95
        if n_verbs == 0:     verb_score = 25
        elif n_verbs <= 3:   verb_score = 25 + n_verbs * 12
        elif n_verbs <= 6:   verb_score = 61 + (n_verbs - 3) * 7
        elif n_verbs <= 10:  verb_score = 82 + (n_verbs - 6) * 3
        else:                verb_score = 94 + min(4, (n_verbs - 10))

        # Penalty per weak phrase
        penalty = len(weak_found) * 8
        score = max(20, min(98, verb_score - penalty))

        return {"score": score, "verbs": verbs_found, "weak": weak_found}

    def _score_quantification(self, text: str) -> dict:
        """
        MOST IMPORTANT differentiator per Jobscan/Resume Worded.
        Count measurable achievements: %, $, numbers, team sizes.
        """
        patterns = [
            r'\b\d+\s*%',                                              # percentages
            r'\$\s*[\d,]+(?:\.\d+)?(?:\s*[MKBmkb](?:illion|illion)?)?\b', # dollar amounts
            r'\b\d+\s*(?:million|billion|thousand|M|K|B)\b',          # large numbers
            r'\b(?:increased?|decreased?|reduced?|improved?|grew?|saved?|generated?|drove?)\w*\s+\w+\s+by\s+\d+', # verb + by X
            r'\b\d+\+?\s*(?:users?|clients?|customers?|projects?|teams?|people|members?|employees?|reports?)\b',
            r'\b\d+x\s*(?:faster|growth|improvement|increase|revenue)\b',
            r'\b(?:top|ranked|#\s*\d+|number\s+\d+)\b',               # rankings
            r'\b\d{4}\s*[-–]\s*(?:\d{4}|present|current)\b',          # date ranges = experience
            r'\bteam\s+of\s+\d+\b',                                    # team size
            r'\b(?:over|more than|nearly|approximately)\s+\d+\b',      # approximate numbers
        ]
        hits = []
        for p in patterns:
            matches = re.findall(p, text, re.IGNORECASE)
            hits.extend([str(m).strip() for m in matches if str(m).strip()])

        n = len(hits)
        # 0=15, 3=55, 6=75, 9=88, 12+=97
        if n == 0:       score = 15
        elif n <= 3:     score = 15 + n * 13
        elif n <= 6:     score = 54 + (n - 3) * 7
        elif n <= 9:     score = 75 + (n - 6) * 4.3
        elif n <= 12:    score = 88 + (n - 9) * 2.7
        else:            score = 96 + min(2, n - 12)

        return {"score": round(min(98, score)), "count": n, "examples": hits[:6]}

    def _score_format(self, text: str) -> dict:
        """
        ATS parsers struggle with: tables, columns, headers/footers,
        special characters, graphics. Check for these issues.
        """
        score = 100
        issues = []

        # Bad special characters that confuse ATS parsers
        bad_chars = re.findall(r'[★✓►▪❖✔☑□■●○◆→←↑↓]', text)
        if len(bad_chars) > 8:
            score -= 20
            issues.append(f"{len(bad_chars)} special characters detected — ATS may misread bullets")

        # Check for likely table/column layout (many pipe chars)
        pipes = text.count('|')
        if pipes > 5:
            score -= 15
            issues.append("Table or column layout detected — some ATS cannot parse tables")

        # Check for very long lines (usually indicates table layout)
        lines = text.split('\n')
        long_lines = [l for l in lines if len(l) > 120]
        if len(long_lines) > 5:
            score -= 10
            issues.append("Unusually long lines detected — may indicate multi-column layout")

        # Email in header/footer repeated = OK
        # Check for garbled OCR text
        non_ascii = len(re.findall(r'[^\x00-\x7F]', text))
        if non_ascii > 50:
            score -= 10
            issues.append("Non-standard characters detected — may cause ATS parsing issues")

        return {"score": max(40, score), "issues": issues}

    def _score_soft_skills(self, lower: str) -> dict:
        """
        Soft skill keywords matter more than most people think.
        LinkedIn's algorithm heavily weights these.
        """
        soft_list = [
            "leadership","communication","collaboration","teamwork",
            "problem solving","critical thinking","analytical","strategic",
            "innovative","adaptable","detail-oriented","self-motivated",
            "proactive","cross-functional","stakeholder","mentoring",
            "coaching","decision making","time management","organized",
            "results-driven","data-driven","customer-focused","team player",
            "fast learner","multi-tasking","deadline","prioritization",
            "interpersonal","presentation","negotiation","creative",
            "project management","planning","execution","accountability",
        ]
        found = [s for s in soft_list if re.search(r'\b' + re.escape(s) + r'\b', lower)]
        n = len(found)

        # 0=20, 2=45, 4=65, 6=80, 8=90, 10+=97
        if n == 0:      score = 20
        elif n <= 2:    score = 20 + n * 12
        elif n <= 4:    score = 44 + (n - 2) * 11
        elif n <= 6:    score = 66 + (n - 4) * 7
        elif n <= 9:    score = 80 + (n - 6) * 3.5
        else:           score = 90 + min(7, (n - 9) * 1.5)

        missing = [s for s in soft_list[:12] if s not in found]
        return {"score": round(min(97, score)), "found": found, "missing": missing}

    def _score_length(self, text: str) -> dict:
        """
        Real ATS tools check word count, bullet density, and section length.
        Ideal: 400-800 words for most roles, up to 1200 for senior.
        """
        words = len(text.split())
        lines = [l.strip() for l in text.split('\n') if l.strip()]
        bullets = len([l for l in lines if re.match(r'^[\-\•\*\–\—]', l)])
        notes = []

        if words < 150:
            score = 20; notes.append(f"Very short ({words} words) — add more detail")
        elif words < 300:
            score = 50; notes.append(f"Short ({words} words) — aim for 400-700 words")
        elif words < 400:
            score = 72; notes.append(f"Slightly short ({words} words) — consider expanding")
        elif words <= 800:
            score = 98; notes.append(f"Ideal length ({words} words)")
        elif words <= 1200:
            score = 85; notes.append(f"Slightly long ({words} words) — consider trimming")
        else:
            score = 60; notes.append(f"Too long ({words} words) — keep to 1-2 pages")

        return {"score": score, "word_count": words, "bullet_count": bullets, "notes": notes}

    def _score_jd_match(self, token_set: set, lower: str, jd: str) -> dict:
        """
        Compare resume against job description.
        Extracts keywords from JD and checks how many appear in resume.
        This is the core of what Jobscan does.
        """
        jd_lower = jd.lower()
        jd_words = re.findall(r'\b[a-z][a-z0-9\+\#\.\-\/]*\b', jd_lower)
        jd_bigrams = [jd_words[i] + " " + jd_words[i+1] for i in range(len(jd_words)-1)]
        jd_all = set(jd_words + jd_bigrams)

        # Remove stopwords and very short words
        jd_meaningful = {t for t in jd_all if t not in STOPWORDS and len(t) > 2}

        # Prioritize: keywords that also appear in universal list
        jd_important = {t for t in jd_meaningful if t in UNIVERSAL_KEYWORDS}
        jd_other = jd_meaningful - jd_important

        # Check matches in resume
        resume_all = token_set | {w for w in lower.split()}

        matched_important = [t for t in jd_important if t in lower or t in token_set]
        matched_other     = [t for t in jd_other     if t in lower or t in token_set]

        all_matched  = sorted(set(matched_important + matched_other))
        all_missing  = sorted((jd_important | jd_other) - set(all_matched))

        total_jd = len(jd_meaningful)
        total_matched = len(all_matched)
        pct = round((total_matched / max(total_jd, 1)) * 100)

        return {
            "match_percentage":  pct,
            "matched_keywords":  all_matched[:22],
            "missing_keywords":  all_missing[:20],
            "jd_word_count":     len(jd.split()),
        }

    # ══════════════════════════════════════════════════════════
    # IMPROVEMENT GENERATOR
    # ══════════════════════════════════════════════════════════

    def _build_improvements(self, scores, r_contact, r_sections,
                             r_keywords, r_writing, r_quant,
                             r_format, r_soft, jd_result) -> list:
        tips = []

        # Contact
        if scores["contact"] < 85:
            missing = r_contact.get("missing", [])
            if missing:
                tips.append({
                    "priority": "HIGH",
                    "category": "Contact Information",
                    "issue": f"Missing: {', '.join(missing)}",
                    "fix": (
                        "Add your LinkedIn profile URL and phone number to the top of your resume. "
                        "Every ATS and recruiter checks these first. "
                        "Format: linkedin.com/in/yourname"
                    )
                })

        # Sections
        if scores["sections"] < 80:
            missing = r_sections.get("missing", [])
            critical = [s for s in missing if s in ["Experience","Education","Skills"]]
            if critical:
                tips.append({
                    "priority": "HIGH",
                    "category": "Resume Structure",
                    "issue": f"Critical sections missing: {', '.join(critical)}",
                    "fix": (
                        "Add clearly labeled section headers. ATS systems scan for exact "
                        "headings like 'Work Experience', 'Education', 'Skills'. "
                        "Without these, your resume may score 0 in those areas."
                    )
                })
            elif missing:
                tips.append({
                    "priority": "MEDIUM",
                    "category": "Resume Structure",
                    "issue": f"Sections not detected: {', '.join(missing[:3])}",
                    "fix": (
                        "Consider adding a Professional Summary section at the top — "
                        "recruiters spend 7 seconds on the first scan. A 3-line summary "
                        "dramatically improves first impressions."
                    )
                })

        # Quantification — most critical per real ATS tools
        if scores["quant"] < 70:
            count = r_quant.get("count", 0)
            tips.append({
                "priority": "HIGH",
                "category": "Quantified Achievements",
                "issue": f"Only {count} measurable results found — this is the #1 differentiator",
                "fix": (
                    "Add numbers to every bullet point where possible. Examples: "
                    "'Increased sales by 32%', 'Led a team of 8 engineers', "
                    "'Reduced load time by 1.2s saving $50K annually', "
                    "'Served 10,000+ users'. Resumes with numbers get 40% more callbacks."
                )
            })

        # Writing quality
        if scores["writing"] < 65:
            weak = r_writing.get("weak", [])
            verbs = r_writing.get("verbs", [])
            if weak:
                tips.append({
                    "priority": "HIGH",
                    "category": "Writing Quality",
                    "issue": f"Weak passive phrases found: \"{weak[0]}\"",
                    "fix": (
                        "Replace passive phrases with strong action verbs. "
                        "❌ 'Responsible for managing team' "
                        "✅ 'Led and mentored a 6-person engineering team'. "
                        "Start every bullet with a past-tense action verb."
                    )
                })
            elif len(verbs) < 4:
                tips.append({
                    "priority": "MEDIUM",
                    "category": "Writing Quality",
                    "issue": f"Only {len(verbs)} action verbs detected",
                    "fix": (
                        "Start each bullet point with a strong action verb. "
                        "Top verbs: Delivered, Architected, Spearheaded, Optimized, "
                        "Drove, Scaled, Reduced, Launched, Led, Generated, Designed."
                    )
                })

        # Keywords
        if scores["keywords"] < 60:
            missing_kw = r_keywords.get("missing", [])[:6]
            tips.append({
                "priority": "HIGH",
                "category": "Hard Skills & Keywords",
                "issue": "Low keyword density — ATS may rank resume low",
                "fix": (
                    f"Add relevant technical skills to your Skills section. "
                    f"Consider including: {', '.join(missing_kw)}. "
                    f"ATS systems rank resumes by keyword match — "
                    f"only resumes with matching keywords reach recruiters."
                )
            })

        # Soft skills
        if scores["soft"] < 50:
            tips.append({
                "priority": "LOW",
                "category": "Soft Skills",
                "issue": "Soft skill keywords are sparse",
                "fix": (
                    "Weave soft skills naturally into experience descriptions. "
                    "E.g., 'Collaborated cross-functionally with 4 teams to deliver...', "
                    "'Mentored 3 junior developers...', 'Led stakeholder presentations...'. "
                    "Don't just list them — show them in context."
                )
            })

        # Format
        if scores["format"] < 70:
            issues = r_format.get("issues", [])
            if issues:
                tips.append({
                    "priority": "MEDIUM",
                    "category": "ATS Format Compatibility",
                    "issue": issues[0],
                    "fix": (
                        "Use a single-column, plain text-friendly layout. Avoid tables, "
                        "text boxes, headers/footers, and graphics — many ATS systems "
                        "cannot parse these and will score your resume lower. "
                        "Use simple bullet points (- or •)."
                    )
                })

        # JD match
        if jd_result and jd_result.get("match_percentage", 100) < 50:
            missing_jd = jd_result.get("missing_keywords", [])[:5]
            tips.append({
                "priority": "HIGH",
                "category": "Job Description Match",
                "issue": f"Only {jd_result['match_percentage']}% keyword match with this job",
                "fix": (
                    f"Mirror the exact language from the job description. "
                    f"Add these missing keywords naturally: {', '.join(missing_jd)}. "
                    f"Jobscan research shows 60%+ match = significantly more interviews. "
                    f"Tailor your resume for each application."
                )
            })

        # Sort: HIGH first, then MEDIUM, then LOW
        order = {"HIGH": 0, "MEDIUM": 1, "LOW": 2}
        tips.sort(key=lambda x: order.get(x["priority"], 3))
        return tips

    # ── Helpers ───────────────────────────────────────────────

    def _tokenize(self, text: str) -> list:
        lower = text.lower()
        words = re.findall(r'\b[a-z][a-z0-9\+\#\.]*\b', lower)
        bigrams = [words[i] + " " + words[i+1] for i in range(len(words)-1)]
        trigrams = [words[i]+" "+words[i+1]+" "+words[i+2] for i in range(len(words)-2)]
        return words + bigrams + trigrams
